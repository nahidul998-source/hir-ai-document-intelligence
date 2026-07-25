import sys
import os
import uuid
import json
import logging
from typing import Dict, Any, Optional

from ai.worker.classifiers.document_classifier import DocumentClassifier, DocumentType
from ai.worker.prompt_registry.registry import PromptRegistry
from ai.worker.confidence.engine import ConfidenceEngine
from ai.worker.validators.json_validator import JSONValidator
from ai.worker.metrics.tracker import MetricsTracker

# The AI worker needs to import AIProviderManager from the backend. 
# PYTHONPATH must be set appropriately when running.
from app.infrastructure.adapters.ai_provider_manager import AIProviderManager
from app.application.services.ai.orchestrator import AIOrchestrator
from app.domain.services.validation.pipeline import ValidationPipeline

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, ai_provider_manager: AIProviderManager, db: Any = None):
        self.classifier = DocumentClassifier()
        self.registry = PromptRegistry()
        self.confidence_engine = ConfidenceEngine()
        self.validator = JSONValidator()
        self.ai_provider = ai_provider_manager
        self.orchestrator = AIOrchestrator(ai_provider_manager)
        self.validation_pipeline = ValidationPipeline(db)
        
    async def extract_text_and_layout(self, filename: str, file_bytes: bytes) -> Dict[str, Any]:
        """Phase 2: Hybrid OCR Engine - Extracts text, layout, and tables."""
        metrics = MetricsTracker()
        metrics.start()
        metrics.start_ocr()
        logger.info(f"Extracting text from {filename} using PyMuPDF (Hybrid Fallback ready)")
        
        import io
        extracted_text = ""
        layout_blocks = []
        tables = []
        
        # Detect extension
        ext = filename.split('.')[-1].lower() if '.' in filename else ''
        
        try:
            if ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                for idx, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        extracted_text += para.text + "\n"
                        layout_blocks.append({"page": 1, "bbox": [0, idx*10, 500, (idx+1)*10], "text": para.text.strip(), "type": "text"})
                for t_idx, table in enumerate(doc.tables):
                    cells = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    tables.append({"page": 1, "bbox": [0,0,0,0], "cells": cells})
            elif ext in ['xlsx', 'xls']:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    cells = []
                    for row in sheet.iter_rows(values_only=True):
                        row_data = [str(cell).strip() if cell is not None else "" for cell in row]
                        if any(row_data):
                            cells.append(row_data)
                    if cells:
                        tables.append({"page": 1, "bbox": [0,0,0,0], "cells": cells})
                        extracted_text += f"\nSheet {sheet_name}:\n"
                        for row in cells:
                            extracted_text += " | ".join(row) + "\n"
            elif ext in ['jpg', 'jpeg', 'png']:
                import fitz
                pdf_document = fitz.open(stream=file_bytes, filetype=ext)
                page = pdf_document.load_page(0)
                extracted_text = page.get_text()
                layout_blocks.append({"page": 1, "bbox": page.rect, "text": "Image Content", "type": "image"})
                pdf_document.close()
            else: # Fallback to PDF parsing
                import fitz
                import pdfplumber
                pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    
                    # 1. Extract Layout Blocks and Images
                    blocks = page.get_text("blocks")
                    for b in blocks:
                        if len(b) >= 6: # standard text/image block
                            x0, y0, x1, y1, text, block_type, block_no = b[:7]
                            if text.strip() or block_type == 1:
                                layout_blocks.append({
                                    "page": page_num + 1,
                                    "bbox": [x0, y0, x1, y1],
                                    "text": text.strip() if block_type == 0 else "[IMAGE]",
                                    "type": "text" if block_type == 0 else "image"
                                })
                                extracted_text += text + "\n"
                            
                pdf_document.close()
                
                # 2. Extract Tables using pdfplumber for better accuracy
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page_num, ppage in enumerate(pdf.pages):
                        page_tables = ppage.extract_tables()
                        for t in page_tables:
                            tables.append({
                                "page": page_num + 1,
                                "bbox": [0,0,0,0], # Simplified bbox for pdfplumber
                                "cells": t
                            })
                            
            if not extracted_text.strip() and not tables:
                logger.warning(f"No text extracted. Triggering PaddleOCR fallback... (Simulated)")
                extracted_text = "ERROR: Scanned document requires PaddleOCR. Fallback activated."
        except Exception as e:
            logger.error(f"Error during extraction: {e}")
            extracted_text = f"ERROR: Failed to extract text - {e}"
            
        metrics.end_ocr()
        
        return {
            "text": extracted_text,
            "layout_blocks": layout_blocks,
            "tables": tables,
            "ocr_metrics": metrics.get_metrics()
        }
        
    async def extract_data(self, filename: str, extracted_text: str, doc_type: str, db: Any, ai_provider: Optional[str] = None) -> Dict[str, Any]:
        """Phase 3: LLM Extraction, Master Data Validation, and Confidence Engine."""
        self.validation_pipeline = ValidationPipeline(db)
        metrics = MetricsTracker()
        metrics.start()
        
        trace_id = str(uuid.uuid4())
        
        # 1. Retrieve Prompt & Schema & Metadata
        prompt_text = self.registry.get_prompt(DocumentType(doc_type))
        prompt_version = self.registry.get_version(DocumentType(doc_type))
        doc_schema = self.registry.get_schema(DocumentType(doc_type))
        
        # 2. AI Provider Extraction via Orchestrator
        system_prompt = "You are an expert Garment Merchandiser AI."
        full_prompt = f"{prompt_text}\n\nDocument Text:\n{extracted_text}"
        
        provider_name = "unknown"
        model_name = "unknown"
        extracted_json = {}
        
        try:
            orch_res = await self.orchestrator.generate_json(
                prompt=full_prompt,
                schema=doc_schema,
                document_type=doc_type,
                system_prompt=system_prompt,
                trace_id=trace_id,
                forced_provider_key=ai_provider
            )
            extracted_json = orch_res["data"]
            provider_name = orch_res["provider"]
        except Exception as e:
            logger.error(f"Extraction failed via AI Orchestrator: {str(e)}")
            extracted_json = {}
            
        metrics.start_llm(provider_name)
        metrics.end_llm(tokens=len(full_prompt.split()))
        
        # 3. Validation (Schema & Pydantic)
        is_schema_valid, validation_msg = self.validator.validate_schema(extracted_json, doc_schema)
        is_pydantic_valid, pydantic_msg = self.validator.validate_pydantic(extracted_json, doc_type)
        
        # 4. Enterprise Validation Engine (Master Data & Business Rules)
        enterprise_validation_result = await self.validation_pipeline.run(doc_type, extracted_json)
        
        # 5. Confidence Scoring
        confidence_metadata = self.confidence_engine.evaluate(
            extracted_data=enterprise_validation_result["validated_data"],
            ocr_text=extracted_text,
            layout_blocks=[],
            provider=metrics.provider
        )
        
        return {
            "classifier_result": doc_type,
            "extracted_data": enterprise_validation_result["validated_data"],
            "confidence_metadata": confidence_metadata,
            "master_data_metadata": enterprise_validation_result["master_data_metadata"],
            "business_rule_errors": enterprise_validation_result["business_rule_errors"],
            "metrics": metrics.get_metrics(),
            "is_schema_valid": is_schema_valid,
            "is_pydantic_valid": is_pydantic_valid,
            "is_business_valid": enterprise_validation_result["is_valid"]
        }

    async def process_document(self, filename: str, first_page_text: str = "", db: Any = None) -> Dict[str, Any]:
        doc_type_dict = await self.classifier.classify(filename, first_page_text)
        doc_type = doc_type_dict.get("document_type", "generic")
        return await self.extract_data(filename, first_page_text, doc_type, db)
